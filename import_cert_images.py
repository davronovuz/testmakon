"""
Docx dan rasmlarni chiqarib, tegishli CertQuestion.image ga saqlash.
Har bir variant uchun rasm → savol mapping aniqlanadi.
"""
import os
import re
import sys
import django

sys.path.insert(0, '/Users/macbookpro/Documents/testmakon')
os.environ['DJANGO_SETTINGS_MODULE'] = 'config.settings'
django.setup()

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from lxml import etree
from certificate.models import CertMock, CertQuestion
from django.core.files.base import ContentFile

DOCX_PATH = '/Users/macbookpro/Desktop/1. Milliy sertifikat [2026].docx'

# Our variant start paragraphs (from successful parse)
# We need to match these to CertMock objects


def get_image_from_paragraph(para, doc):
    """Extract image bytes from a paragraph's inline drawing."""
    NS_W = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
    NS_R = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'
    NS_DRAW = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
    NS_WP = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
    NS_PIC = '{http://schemas.openxmlformats.org/drawingml/2006/picture}'

    drawings = para._element.findall(f'.//{NS_W}drawing')
    images = []
    for drawing in drawings:
        # Find blip (image reference)
        blips = drawing.findall(f'.//{DRAW_NS}blip')
        if not blips:
            # Try alternative namespace
            blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
        for blip in blips:
            embed = blip.get(f'{NS_R}embed')
            if not embed:
                embed = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
            if embed:
                try:
                    image_part = doc.part.related_parts[embed]
                    images.append({
                        'bytes': image_part.blob,
                        'content_type': image_part.content_type,
                        'ext': image_part.content_type.split('/')[-1].replace('jpeg', 'jpg'),
                    })
                except KeyError:
                    pass
    return images


DRAW_NS = '{http://schemas.openxmlformats.org/drawingml/2006/main}'
R_NS = '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}'


def extract_images_simple(doc):
    """Simpler approach: find all images in order and map to paragraphs."""
    paras = doc.paragraphs
    para_images = {}

    for i, para in enumerate(paras):
        el = para._element
        # Find all drawing elements
        drawings = el.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
        if not drawings:
            continue

        for drawing in drawings:
            blips = drawing.findall(f'.//{DRAW_NS}blip')
            for blip in blips:
                embed = blip.get(f'{R_NS}embed')
                if embed and embed in doc.part.related_parts:
                    image_part = doc.part.related_parts[embed]
                    ext = image_part.content_type.split('/')[-1].replace('jpeg', 'jpg')
                    if i not in para_images:
                        para_images[i] = []
                    para_images[i].append({
                        'bytes': image_part.blob,
                        'ext': ext,
                    })

    return para_images


def find_variant_ranges(paras):
    """Same logic as parser - find variant starts with highlighted answers."""
    from docx.enum.text import WD_COLOR_INDEX
    HIGHLIGHT_COLORS = {WD_COLOR_INDEX.TURQUOISE, WD_COLOR_INDEX.BRIGHT_GREEN}

    candidates = []
    for i, para in enumerate(paras):
        text = para.text.strip()
        if re.match(r'^1\.\s*[\[\(]\s*[\d,\.]+\s*ball', text, re.I):
            candidates.append(i)
        elif re.match(r'^1\.\s+[A-ZА-ЯTQ]', text) and not re.match(r'^1\.\s*(BaCl|X[\s,]|A\s+(va|modda)|Sxema|Dastlabki|Jarayon|Ushbu|Toʻyingan|Glyukoza)', text):
            candidates.append(i)

    real_starts = []
    for ci, c in enumerate(candidates):
        end = candidates[ci + 1] if ci + 1 < len(candidates) else len(paras)
        hl_count = 0
        for j in range(c, min(end, c + 500)):
            for run in paras[j].runs:
                if run.font.highlight_color in HIGHLIGHT_COLORS:
                    m = re.match(r'^[A-D]\)', run.text.strip())
                    if m:
                        hl_count += 1
                        break
        if hl_count >= 15:
            real_starts.append(c)

    return real_starts


def find_question_for_image(paras, img_para_idx, variant_start, variant_end):
    """Given an image paragraph, find which question number it belongs to."""
    # Look at the image paragraph itself first
    text = paras[img_para_idx].text.strip()
    m = re.match(r'^(\d+)\.', text)
    if m:
        return int(m.group(1))

    # Look backwards for a question number
    for j in range(img_para_idx - 1, max(variant_start - 1, img_para_idx - 5), -1):
        t = paras[j].text.strip()
        m = re.match(r'^(\d+)\.\s', t)
        if m:
            num = int(m.group(1))
            if 1 <= num <= 45:
                return num

    # Look forward
    for j in range(img_para_idx + 1, min(variant_end, img_para_idx + 3)):
        t = paras[j].text.strip()
        m = re.match(r'^(\d+)\.\s', t)
        if m:
            num = int(m.group(1))
            # Image before question = probably belongs to previous question or this one
            return max(1, num - 1)

    return None


def main():
    doc = Document(DOCX_PATH)
    paras = doc.paragraphs

    # Get all images by paragraph
    para_images = extract_images_simple(doc)
    print(f"Jami {len(para_images)} rasmli paragraf topildi")

    # Get variant ranges
    variant_starts = find_variant_ranges(paras)
    print(f"Variant starts: {variant_starts}")

    # Get all mocks from DB (ordered by variant number)
    mocks = list(CertMock.objects.filter(
        cert_subject__subject__slug='kimyo'
    ).order_by('slug'))
    print(f"DB da {len(mocks)} mock mavjud")

    if len(mocks) != len(variant_starts):
        print(f"WARNING: mock soni ({len(mocks)}) != variant soni ({len(variant_starts)})")

    saved_count = 0
    for vi, vs in enumerate(variant_starts):
        if vi >= len(mocks):
            break

        mock = mocks[vi]
        ve = variant_starts[vi + 1] if vi + 1 < len(variant_starts) else len(paras)

        # Find images in this variant range
        for img_para, images in para_images.items():
            if img_para < vs or img_para >= ve:
                continue

            q_num = find_question_for_image(paras, img_para, vs, ve)
            if q_num is None or q_num > 32:
                continue

            # Find corresponding CertQuestion
            try:
                question = CertQuestion.objects.get(mock=mock, number=q_num)
            except CertQuestion.DoesNotExist:
                print(f"  Savol topilmadi: {mock.title} Q#{q_num}")
                continue

            if question.image:
                # Already has image
                continue

            # Save first image
            img_data = images[0]
            filename = f"cert_q_{mock.slug}_{q_num}.{img_data['ext']}"
            question.image.save(filename, ContentFile(img_data['bytes']), save=True)
            saved_count += 1
            print(f"  ✓ {mock.title} Q#{q_num} → {filename}")

    print(f"\nJami {saved_count} ta rasm saqlandi")


if __name__ == '__main__':
    main()
